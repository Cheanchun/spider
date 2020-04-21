# encoding: utf-8
import json
import os
import random
import re
import time
import traceback

import execjs
import numpy as np
import pandas as pd
import redis
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from pandas import DataFrame
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By  # 选择方法
from selenium.webdriver.support import expected_conditions as EC  # 等待条件
from selenium.webdriver.support.ui import WebDriverWait  # 等待

from helper.config import BaseConfig  # 相关配置


# 将信息批量导入excel表格
class TableHandle:
    def __init__(self, old_table, new_table, concat_columns: list):
        if os.path.isfile(old_table) and os.path.isfile(new_table):
            self.old_table = old_table
            self.new_table = new_table
            self.concat_columns = concat_columns
        else:
            raise FileNotFoundError("FileNotFoundError:文件不存在!")

    def concat_table(self):
        old_data = pd.read_excel(self.old_table)
        new_data = pd.read_excel(self.new_table)[["严重违法失信企业名称", "证件信息(统一社会信用代码)", "生效日期"]]
        data = pd.concat([old_data, pd.DataFrame(np.array(new_data), columns=self.concat_columns)], ignore_index=True)
        data.to_excel("./temp.xlsx")
        data = DataFrame(data)
        print(data[self.concat_columns])


def exec_js2(resp):
    js = "function dd(){var json=" + resp + ";return json.map( function(item){ return String." \
                                            "fromCharCode(item);}).join('');}"
    ctx = execjs.compile(js)
    return ctx.call("dd")


class CompanyInfo(BaseConfig):

    def __init__(self):
        # self.option = webdriver.ChromeOptions()
        # self.option.add_argument(f'--proxy-server=http://{random.choice(self.PROXIES).get("http")}')  # 设置代理
        # # self.option.add_argument(f'--proxy-server=http://ip:port')  # 设置代理
        # # self.option.add_argument("--headless")
        # # self.option.headless = True
        # self.browser = webdriver.Chrome(chrome_options=self.option)  #无头模式启用获取的cookies 不能用于requests请求
        # # self.browser = webdriver.Chrome()
        # self.browser.maximize_window()  # 窗口最大化
        # time.sleep(1)
        # 开发阶段在init中选择一个代理
        self.ip_list = []
        self.session = requests.Session()
        # self.session.proxies = self.proxy  # 如需代理,需要和selenium同时设置
        self.browser = None
        self.r = redis.Redis(host=self.HOST, port=self.PORT, db=self.DB_INDEX, password=self.PASSWORD)
        self.info_base_url = "http://www.gsxt.gov.cn"
        self.brief_msg = None
        self.keyword = None
        self.headers = None
        self.success = None
        self.gt = None
        self.challenge = None
        self.validate = None
        self.info_url = None
        self.query_page = "http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html"
        self.cookies = {}
        self.headers = {
            'Host': 'www.gsxt.gov.cn',
            'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/'
                          '537.36 (KHTML, like Gecko) Chrome/72.0.3610.2 Safari/537.36',
            'Refer': 'http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html',
            'Referer': 'http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html',
            'Origin': 'http://www.gsxt.gov.cn',
            'Accept': 'application/json, text/javascript, */*; q=0.01',
        }

    # 获取 redis数据库中 order_company 有序集合 的一条数据
    def get_redis_data(self):
        '''
        该方法主要是获取redis中数据,取出失信企业的注册码,如果没有注册码,则匹
        配公告内容中的公司名称,使用注册码查询相对文字较为稳定,不容易ban掉.
        :return:notice_url  keyword
        '''
        key = self.r.zrange("order_company", start=0, end=0, desc=False, withscores=True, score_cast_func=int)[0][
            0].decode("utf-8")
        self.brief_msg = eval(key)
        print(self.brief_msg)
        notice_url = self.brief_msg.get("notice_url")
        keyword = self.brief_msg.get("reg_num")
        if not keyword:
            keyword = re.search(pattern="^[^关于].+(服务部|公司|店|网吧|厅|厂|馆|中心|药房|商行|经营部|服务站|招待所|养殖场|研究所)",
                                string=self.brief_msg.get("noticeTitle"))
            if keyword:
                keyword = keyword.group()
                return notice_url, keyword
            keyword = re.search(pattern="关于(.+(服务部|公司|店|网吧|厅|厂|馆|中心|药房|商行|经营部|服务站|招待所|养殖场|研究所))",
                                string=self.brief_msg.get("noticeTitle"))
            if keyword:
                keyword = keyword.group(1)
                return notice_url, keyword
            if not keyword:
                with open("./log/keywordfail.txt", "a+", encoding="utf-8") as fp:
                    fp.write(self.brief_msg.get("noticeTitle") + "\n")
                print("匹配失败", self.brief_msg.get("noticeTitle"))
                self.r.zincrby("order_company", value=str(self.brief_msg), amount=2)
                raise AttributeError("AttributeError:未匹配到公司名称")
        else:
            return notice_url, keyword

    def get_ip(self):
        ip_url = "http://api.xdaili.cn/xdaili-api//privateProxy/applyStaticProxy?spiderId=bcf4b1d29b374d10a6a1d14fd08e516f&returnType=1&count=1"
        if self.ip_list:
            return self.ip_list.pop()
        else:
            self.ip_list = requests.get(url=ip_url).text.split()
            print(self.ip_list)
            return self.ip_list.pop()

    def open_chrome(self):
        self.option = webdriver.ChromeOptions()
        self.browser = webdriver.Chrome()
        self.browser.maximize_window()  # 窗口最大化

    # 获取cookies
    def get_cookies(self):
        '''
        使用selenium 动态获取cookie中__jsluid 和 __jsl_clearance 的值,目标
        网站主要验证cookies中这两个字段信息,__jsl_clearance js代码动态生成
        '''

        print('[*] 1. 第一次访问或者cookies失效，状态码为521，正在通过selenium获取...')
        self.browser.get(self.query_page)
        wait_gonggao_page = WebDriverWait(self.browser, 20)
        for i in range(3):
            flag = False
            try:
                flag = wait_gonggao_page.until(
                    EC.text_to_be_present_in_element((By.XPATH, "//div[@class='page_left']/a[5]/div[@id='tab4']"),
                                                     '严重违法失信企业名单公告'))
            except TimeoutException:
                if i == 2:
                    raise EOFError("cookies页面加载失败")
                pass
            if flag:
                break
        self.cookies = {cookie.get("name"): cookie.get("value") for cookie in self.browser.get_cookies()}
        print(self.cookies)

    def change_proxy(self):
        '''
        ip代理切换,requests和selenium需要更换相同的ip,否则静态请求会导致失败!
        :return:None
        '''
        self.browser.quit()  # 退出当前浏览器
        proxy = self.get_ip()
        self.option.add_argument(f"--proxy-server=http://{proxy.get('http')}")  # 设置浏览器代理
        self.browser = webdriver.Chrome(self.option)
        self.session.proxies = proxy  # 设置 session 静态请求代理
        print(proxy)

    # 获取验证码信息
    def get_image_gif(self):
        url = f'http://www.gsxt.gov.cn/SearchItemCaptcha?t={time.time()*1000}'
        resj = self.session.get(url, headers=self.headers, cookies=self.cookies)
        for_test_error = resj.text  # 用于错误捕获打印
        try:
            resj = resj.json()
            url = "http://www.gsxt.gov.cn/corp-query-custom-geetest-image.gif?v="
            local_time = time.localtime(time.time())
            url = url + str(local_time.tm_min + local_time.tm_sec)
            resp = self.session.get(url, headers=self.headers, cookies=self.cookies)
            self.success = resj['success']
            self.gt = resj['gt']
            self.challenge = resj['challenge']
            try:
                aaa = exec_js2(resp.text)
                matchObj = re.search('location_info = (\d+);', aaa)
            except execjs._external_runtime.ProcessExitedWithNonZeroStatus:
                print("重新执行get_image_gif()")
                self.get_image_gif()
            if matchObj:
                return matchObj.group(1)
            else:
                raise RuntimeError("RuntimeError:没有找到location_info")
        except json.JSONDecodeError as e:
            print("[!]image_gif获取失败,重新获取")
            print(e)
            print(for_test_error)
            self.get_image_gif()

    # 获取token
    def get_token(self):
        url = "http://www.gsxt.gov.cn/corp-query-geetest-validate-input.html?token=" + self.get_image_gif()
        resp = self.session.get(url, headers=self.headers, cookies=self.cookies)
        aaa = exec_js2(resp.text)
        matchObj = re.search('value: (\d+)}', aaa)
        if matchObj:
            location_info = matchObj.group(1)
            self.token = str(int(location_info) ^ 536870911)
            print('[*] 3. 获取token: ' + self.token)
        else:
            raise RuntimeError("RuntimeError:没有找到location_info")

    # 调用JiYan 打码平台进行验证码验证
    def get_validate(self):
        url = f'http://jiyanapi.c2567.com/shibie?gt={self.gt}&challenge={self.challenge}&referer=http://www.gsxt.gov.cn&user={self.USER}&pass={self.PASSWD}&return=json&format=utf8'
        if not self.success:
            url = url + '&model=1'
        res = self.session.get(url=url).text
        resj = json.loads(res)
        if resj['status'] == 'ok':
            print('[*] 4. 对接极验打码平台，获取validate成功')
            print(resj)
            self.validate = resj['validate']
        else:
            raise RuntimeError("RuntimeError:打码平台", resj['msg'])

    def query(self, keyword):
        '''
        以公司注册号或者公司名称查询,查询成功返回True,失败返回False
        :param keyword: 关键字为两种,一种是注册码,另一种为公司名称
        :return: True  False 查询成功标志
        '''

        url = "http://www.gsxt.gov.cn/corp-query-search-1.html"
        post_data = {
            'tab': self.SEARCH_TYPE,  # ill_tab 严重 execp_tab异常 ent_tab 所有
            'province': '',
            'geetest_challenge': self.challenge,
            'geetest_validate': self.validate,
            'geetest_seccode': self.validate + '|jordan',
            'token': self.token,
            'searchword': keyword
        }
        resp = self.session.post(url=url, headers=self.headers, data=post_data, cookies=self.cookies,
                                 allow_redirects=False)
        while resp.status_code != 200 or not resp.text:  # 判断是否重定向和空页面返回
            print(resp.status_code, resp.text)
            print("[!]需要更换代理IP")
            self.change_proxy()
            print("[!]更换cookies")
            self.get_cookies()
            resp = self.session.post(url=url, headers=self.headers, data=post_data, cookies=self.cookies,
                                     allow_redirects=False)
            time.sleep(random.randint(2, 4))
        query_result = re.search('<span class="search_result_span1">(\d+)</span>', resp.text)
        if query_result:
            query_result = query_result.group(1)
        else:
            print("查询失败")
            print(resp.text)
            return False
        if query_result is "1":
            try:
                res = re.search(pattern="target='_blank' href='(.+)'", string=resp.text).group(1)  # 匹配公司详情url
            except AttributeError:
                print("[!]详情url获取失败!")
                return False
            self.info_url = self.info_base_url + res
            print(self.info_url)
            return True
        elif query_result is "0":
            print("[!]未查询到相关信息", resp)  # 查询不到相关信息
            print(resp.text)
            return False
        elif query_result:
            return False
        else:
            self.query(keyword)

    # 获取公司的详细信息页面,公告信息页面,进行处理
    def get_detail_img(self, notice_url):
        '''
        使用selenium访问 url,获取公告页面和企业详情页面,并截图获取文本信息
        将图片信息作为临时文件保存在本地(./img_temp/),文本信息作为返回值
        '''
        time.sleep(random.randint(2, 4))
        self.browser.get(notice_url)
        # 处理公告页面
        notice_wait = WebDriverWait(self.browser, 10)

        flag_notice = False
        try:
            flag_notice = notice_wait.until(
                EC.presence_of_element_located((By.ID, "datefornow"), ))
        except Exception as e:
            print("[!]页面请求失败", e)
            refresh_num = 0
            while not flag_notice and refresh_num < 3:
                print("[!]页面请求失败,重新刷新页面")
                self.browser.refresh()
                try:
                    flag_notice = notice_wait.until(
                        EC.presence_of_element_located((By.ID, "datefornow"), ))
                except Exception as e:
                    print(f"1刷新次数{refresh_num}", e)
                    refresh_num += 1

        self.notice_text = self.browser.find_element(By.XPATH,
                                                     "//div[@class='mainContent']/div[@class='contentfordetail']").text

        # 字符串 按照 \n 切割操作 取出 日期 last
        temp_data = self.notice_text.split()
        if len(temp_data):
            self.date = temp_data[-1]
        else:
            print("[!]生效时间获取失败")
            print(temp_data)
            raise AttributeError("AttributeError:生效时间获取失败!")
        js = self.browser.find_element_by_xpath("//div[@class='page']/div")
        self.browser.execute_script("arguments[0].scrollIntoView(true);", js)
        self.browser.save_screenshot("./img_temp/notice_temp.png")
        # 处理详情页面
        time.sleep(random.randint(2, 3))
        self.browser.get(self.info_url)
        wait_info_page = WebDriverWait(self.browser, 10)
        flag = False
        try:
            flag = wait_info_page.until(
                EC.text_to_be_present_in_element(
                    (By.XPATH, "//div[@id='primaryInfo']//div[@class='classify']"),
                    '营业执照信息'))
        except Exception as e:
            print("[!]页面请求失败", e)
            refresh_num = 0
            while flag is not True and refresh_num < 3:
                print("[!]页面请求失败,重新刷新页面")
                self.browser.refresh()
                try:
                    flag = wait_info_page.until(
                        EC.text_to_be_present_in_element(
                            (By.XPATH, "//div[@id='primaryInfo']//div[@class='classify']"),
                            '营业执照信息'))
                except Exception as e:
                    print(f"2刷新次数{refresh_num}", e)
                    refresh_num += 1

        js = self.browser.find_element_by_id("btn_send_pdf")
        self.browser.execute_script("arguments[0].scrollIntoView(true);", js)
        self.browser.save_screenshot("./img_temp/info_temp.png")
        self.title = self.browser.find_element(By.XPATH, "//div[@class='overview']/dl[2]/dd[@class='result']").text
        self.reg_num = self.browser.find_element(By.XPATH, "//div[@class='overview']/dl[1]/dd[@class='result']").text
        data_for_exc = [self.title, self.reg_num, self.date]  # 公司名称,注册号,生效日期
        data_for_doc = {
            "notice_text": self.notice_text,  # 公告文本
            "reg_num": str(self.reg_num),
        }
        return data_for_exc, data_for_doc

    # 保存为excel和word
    def save_excel(self, data, exc_path="./失信名单客户登记表.xlsx"):
        '''

        :param data:
        :param exc_path:
        :return:
        '''
        insert_columns = ["严重违法失信企业名称", "证件信息(统一社会信用代码)", "生效日期"]
        df = pd.read_excel(exc_path, index_col=0, sort=False)
        update_df = pd.DataFrame([data], columns=insert_columns)
        new_data = df.append(update_df, ignore_index=True)
        new_data.to_excel(exc_path)
        print("保存到Excel文档成功")

    def save_doc(self, data, doc_path="./失信企业信息收集.docx"):
        '''

        :param data:
        :param doc_path:
        :return:
        '''
        document = Document(docx=doc_path)
        text = data.get("notice_text")
        str_list = text.split("\n")
        title = str_list[0]
        sec_title = str_list[1]
        content = "      ".join(i.strip() + "\n" for i in str_list[2:-2])
        last = "\n".join(i.strip() for i in str_list[-2:])

        # title居中
        doc_title = document.add_paragraph()
        run_title = doc_title.add_run(title)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_title.bold = True
        # 小标题居中
        doc_title = document.add_paragraph(sec_title)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 段落
        doc_title = document.add_paragraph(content)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 落款居右
        doc_title = document.add_paragraph(last)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # notice_img
        img_notice = "./img_temp/notice_temp.png"
        document.add_picture(img_notice, width=Inches(5.7))

        # 注册码
        reg_num_info = f"注册码:{data.get('reg_num')}"
        document.add_paragraph(reg_num_info)

        # info_img
        document.add_paragraph("\n")
        img_info = "./img_temp/info_temp.png"
        document.add_picture(img_info, width=Inches(5.7))
        document.add_paragraph("\n")
        document.save("./失信企业信息收集.docx")
        self.r.zincrby("order_company", value=str(self.brief_msg), amount=1)  # 一个redis信息处理完成,将score分数改为1,防止重复使用
        print("保存到word文档成功")

    def __del__(self):
        # 关闭浏览器
        self.browser.quit()


def main():
    # r = redis.Redis(host=BaseConfig.HOST, port=BaseConfig.PORT, db=0, password=BaseConfig.PASSWORD)  # 连接redis
    company_info = CompanyInfo()
    company_info.open_chrome()
    while True:
        counts = company_info.r.zcount("order_company", 0, 0)
        if counts:
            print(f"[!]待处理数据量:{counts}")
            company_info.get_cookies()
            try:
                for i in range(counts):
                    notice_url, query_word = company_info.get_redis_data()
                    company_info.get_token()
                    company_info.get_validate()
                    if query_word and company_info.query(query_word):
                        exc_data, doc_data = company_info.get_detail_img(notice_url)
                        company_info.save_doc(data=doc_data)
                        company_info.save_excel(data=exc_data)
                    elif query_word:
                        data = [query_word, query_word, "未查询到详情或查询不准确"]
                        company_info.save_excel(data=data)
                        with open("./log/out_of_query.txt", 'a+', encoding="utf-8") as fp:
                            fp.write(str(query_word) + "\n")
                            company_info.r.zincrby("order_company", value=str(company_info.brief_msg),
                                                   amount=2)  # 未查询到详细信息的企业,跳过,标记为2
                    print(f"[√]本次运行已处理{i+1}条")
                company_info.__del__()
            except (Exception, BaseException):
                company_info.__del__()
                # traceback.print_exc(file=open("./log/error.log", "a+", encoding="utf-8"),)
                with open("./log/error.log", "a+", encoding="utf-8") as fp:
                    fp.write(time.strftime("%Y-%m-%d %H:%M:%S",
                                           time.localtime(time.time())) + ":" + traceback.format_exc() + "\n")
                traceback.format_exc()

        else:
            print("[!]暂无待处理数据")
            time.sleep(60)


if __name__ == '__main__':
    main()
