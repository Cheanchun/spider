# encoding: utf-8
import json
import logging.config
import time
import traceback
from datetime import datetime

import execjs
import os
import pandas as pd
import random
import re
import redis
import requests
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from lxml import etree
from pyvirtualdisplay import Display
from selenium import webdriver
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By  # 选择方法
from selenium.webdriver.support import expected_conditions as EC  # 等待条件
from selenium.webdriver.support.ui import WebDriverWait  # 等待

from helper.config import BaseConfig  # 相关配置

logging.basicConfig(level=logging.DEBUG,
                    format="%(asctime)s %(name)s %(levelname)s %(message)s",
                    datefmt='%Y-%m-%d  %H:%M:%S %a',
                    filename=BaseConfig.HOME_PATH + "/log/error.log"  # 有了filename参数就不会直接输出显示到控制台，而是直接写入文件
                    )

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOG_DIR = os.path.join(BASE_DIR, "logs/img/")
if not os.path.exists(LOG_DIR):
    os.makedirs(LOG_DIR)
LOG_FILE = datetime.now().strftime("%Y-%m-%d") + "gsxt.log"
LOGGING = {
    "version": 1,
    "disable_existing_loggers": False,
    "formatters": {
        "simple": {
            'format': '%(asctime)s [%(name)s:%(lineno)d] [%(levelname)s]- %(message)s'
        },
        'standard': {
            'format': '%(asctime)s [%(threadName)s:%(thread)d] [%(name)s:%(lineno)d] [%(levelname)s]- %(message)s'
        },
    },

    "handlers": {
        "console": {
            "class": "logging.StreamHandler",
            "level": "DEBUG",
            "formatter": "simple",
            "stream": "ext://sys.stdout"
        },

        "default": {
            "class": "logging.handlers.RotatingFileHandler",
            "level": "INFO",
            "formatter": "simple",
            "filename": os.path.join(LOG_DIR, LOG_FILE),
            'mode': 'w+',
            "maxBytes": 1024 * 1024 * 5,  # 5 MB
            "backupCount": 10,
            "encoding": "utf8"
        },
    },

    "root": {
        'handlers': ['default'],
        'level': "INFO",
        'propagate': False
    }
}

logging.config.dictConfig(LOGGING)
gsxt_logger = logging.getLogger(__file__)


# 当元素出现
def is_visible(driver, locator, timeout=10):
    try:
        WebDriverWait(driver, timeout).until(EC.visibility_of_element_located((By.XPATH, locator)))
        return True
    except TimeoutException:
        return False


# 当元素消失
def is_not_visible(driver, locator, timeout=10):
    try:
        WebDriverWait(driver, timeout).until_not(EC.visibility_of_element_located((By.XPATH, locator)))
        return True
    except TimeoutException:
        return False


def exec_js2(resp):
    js = "function dd(){var json=" + resp + ";return json.map( function(item){ return String." \
                                            "fromCharCode(item);}).join('');}"
    ctx = execjs.compile(js)
    return ctx.call("dd")


class CompanyInfo(BaseConfig):
    def __init__(self):
        self.display = Display(visible=0, size=(1280, 800))  # 创建虚拟展示界面
        self.display.start()  # 开始
        self.session = requests.Session()
        self.browser = None
        self.r = redis.Redis(host=self.HOST, port=self.PORT, db=self.DB_INDEX, password=self.PASSWORD)
        self.info_base_url = "http://www.gsxt.gov.cn"
        self.brief_msg = None
        self.keyword = None
        self.headers = None
        self.success = None
        self.gt = None
        self.challenge = None
        self.validate = None
        self.info_url = None
        self.ip_list = []
        self.query_page = "http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html"
        self.cookies = {}
        self.headers = {
            'Host': 'www.gsxt.gov.cn',
            # 129
            'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/71.0.3578.98 Safari/537.36',
            # 87
            # 'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/73.0.3683.75 Safari/537.36', # 87
            # windows
            # 'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3610.2 Safari/537.36', # windows
            'Refer': 'http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html',
            'Referer': 'http://www.gsxt.gov.cn/corp-query-entprise-info-xxgg-100000.html',
            'Origin': 'http://www.gsxt.gov.cn',
            'Accept': 'application/json, text/javascript, */*; q=0.01',
        }

    # 获取 redis数据库中,有序集合的一条数据
    def get_redis_data(self):
        '''
        该方法主要是获取redis中数据,取出失信企业的注册码,如果没有注册码,则匹
        配公告内容中的公司名称,使用注册码查询相对文字较为稳定,不容易ban掉.
        :return:notice_url  keyword
        '''
        key = self.r.zrange(self.REDIS_SET_NAME, start=0, end=0, desc=False, withscores=True, score_cast_func=int)[0][
            0].decode("utf-8")
        self.brief_msg = eval(key)
        notice_url = self.brief_msg.get("notice_url")
        reg_num = self.brief_msg.get("reg_num")
        company_name = self.brief_msg.get("company_name")
        return notice_url, reg_num, company_name

    def open_chrome(self):
        options = webdriver.ChromeOptions()
        options.add_argument('--disable-gpu')
        options.add_argument('--no-sandbox')
        self.browser = webdriver.Chrome(chrome_options=options)
        self.browser.maximize_window()

    # 获取cookies
    def get_cookies(self):
        '''
        使用selenium 动态获取cookie中__jsluid 和 __jsl_clearance 的值,目标
        网站主要验证cookies中这两个字段信息,__jsl_clearance js代码动态生成
        '''

        gsxt_logger.info('[*] 1. 第一次访问或者cookies失效，状态码为521，正在通过selenium获取...')
        self.browser.get(self.query_page)
        wait_gonggao_page = WebDriverWait(self.browser, 20)
        for i in range(3):  # 出现知道创宇监测页面/和ip检测页面   --->time.sleep()减缓访问速度
            flag = False
            try:
                flag = wait_gonggao_page.until(
                    EC.text_to_be_present_in_element((By.XPATH, "//div[@class='page_left']/a[5]/div[@id='tab4']"),
                                                     '严重违法失信企业名单公告'))
            except TimeoutException:
                logging.info("刷新cookie获取页面：{}".format(i))
                time.sleep(random.randint(2, 5))
                self.browser.get(self.query_page)
            if flag:
                break
            elif i == 2:
                gsxt_logger.info("[!]cookies 页面刷新失败!")
                raise EOFError("[!]cookies 页面刷新失败!")
            else:
                pass
        self.cookies = {cookie.get("name"): cookie.get("value") for cookie in self.browser.get_cookies()}

    # 获取验证码信息
    def get_image_gif(self):
        url = f'http://www.gsxt.gov.cn/SearchItemCaptcha?t={time.time()*1000}'
        resj = self.session.get(url, headers=self.headers, cookies=self.cookies)
        for_test_error = resj.text  # 用于错误捕获打印
        for i in range(3):
            try:
                resj = resj.json()
                url = "http://www.gsxt.gov.cn/corp-query-custom-geetest-image.gif?v="
                local_time = time.localtime(time.time())
                url = url + str(local_time.tm_min + local_time.tm_sec)
                time.sleep(random.randint(1, 2))
                resp = self.session.get(url, headers=self.headers, cookies=self.cookies)
                self.success = resj['success']
                self.gt = resj['gt']
                self.challenge = resj['challenge']
                aaa = exec_js2(resp.text)
                matchObj = re.search('location_info = (\d+);', aaa)
                if matchObj:
                    return matchObj.group(1)
                else:
                    gsxt_logger.info("RuntimeError:没有找到location_info")
                    raise EOFError("RuntimeError:没有找到location_info")
            except json.JSONDecodeError:
                gsxt_logger.info(f"[!]image_gif获取失败,重新获取。{for_test_error}")
                time.sleep(2)
        gsxt_logger.info("image_gif获取失败,重新获取")
        raise EOFError("image_gif获取失败,重新获取")

    # 获取token
    def get_token(self):
        url = "http://www.gsxt.gov.cn/corp-query-geetest-validate-input.html?token=" + self.get_image_gif()
        resp = self.session.get(url, headers=self.headers, cookies=self.cookies)
        time.sleep(random.randint(1, 2))
        aaa = exec_js2(resp.text)
        matchObj = re.search('value: (\d+)}', aaa)
        if matchObj:
            location_info = matchObj.group(1)
            self.token = str(int(location_info) ^ 536870911)
        else:
            gsxt_logger.info("[!]token获取失败,等待重新获取!")
            time.sleep(random.randint(3, 5))
            self.get_token()

    # 调用JiYan 打码平台进行验证码验证
    def get_validate(self):
        try:
            url = f"http://jiyanapi.c2567.com/shibie?gt={self.gt}&challenge={self.challenge}&referer=http://www.gsxt.gov.cn/&user={self.USER}&pass={self.PASSWD}&return=json&format=utf8"
            if not self.success:
                url = url + '&model=3'
            res = self.session.get(url=url).text
            resj = json.loads(res)  # todo 有一定几率报错
            if resj['status'] == 'ok':
                print('[*] 4. 对接极验打码平台，获取validate成功')
                print(resj)
                self.validate = resj['validate']
                # 如果是滑动验证 则需要使用极验返回的challenge,如果是点选,则使用原来的challenge
                self.challenge = resj["challenge"]
            else:
                print(f"[!]极验打码平台验证失败!,等待重新验证!{resj}")
                gsxt_logger.info(f"[!]极验打码平台验证失败!,等待重新验证!{resj}")
                time.sleep(random.randint(2, 3))
                self.get_token()
                self.get_validate()
        except (Exception, BaseException) as e:
            gsxt_logger.info(f"验证出错!重新验证{e}")
            self.get_validate()

    def get_earliest_time(self, html):
        '''
        :return:date
        '''
        res = re.search(pattern='IllInfoUrl = "(.+?)";', string=html)
        if res:
            res = res.group(1)
        else:
            time.sleep(random.randint(5, 15))
            gsxt_logger.info("获取列入失信列表时间,API_url匹配失败")
            raise NameError("API_url匹配失败")
        url = u"http://www.gsxt.gov.cn" + res
        # print(url)
        res = self.session.get(url=url, headers=self.headers, cookies=self.cookies, allow_redirects=False)
        datas = res.json().get("data")
        dates = [int(data.get("abntime")) for data in datas]
        if dates:
            date = min(dates)
            return time.strftime("%Y年%m月%d日", time.localtime(date / 1000))
        else:
            return "None"

    def query(self, keyword):
        '''
        以公司注册号或者公司名称查询,查询成功返回True,失败返回False
        :param keyword: 关键字为两种,一种是注册码,另一种为公司名称
        :return: True  False 查询成功标志
        '''
        self.get_token()
        self.get_validate()
        url = u"http://www.gsxt.gov.cn/corp-query-search-1.html"
        post_data = {
            u'tab': self.SEARCH_TYPE,  # ill_tab 严重 execp_tab异常 ent_tab 所有
            u'province': '',
            u'geetest_challenge': self.challenge,
            u'geetest_validate': self.validate,
            u'geetest_seccode': self.validate + u'|jordan',
            u'token': self.token,
            u'searchword': keyword
        }
        time.sleep(random.randint(4, 9))
        resp = self.session.post(url=url, headers=self.headers, data=post_data, cookies=self.cookies,
                                 allow_redirects=False)
        # print(resp.headers)
        repeat_num = 0
        while resp.status_code != 200 or not resp.text:  # 判断是否重定向和空页面返回
            gsxt_logger.info(f"[!]查询页面错误{resp.status_code}")
            repeat_num += 1
            if repeat_num > 3:
                return False
            time.sleep(random.randint(20, 30))
            if resp.status_code == 521:
                gsxt_logger.info(u"查询cookie失效,重新获取cookies")
                self.get_cookies()
            self.get_token()
            self.get_validate()
            resp = self.session.post(url=url, headers=self.headers, data=post_data, cookies=self.cookies,
                                     allow_redirects=False)
        query_result = re.search(u'<span class="search_result_span1">(\d+)</span>', resp.text)
        if query_result:
            query_result = query_result.group(1)
        else:
            gsxt_logger.info(f"查询失败,score:2,{keyword}")
            # print resp.text
            return False
        if query_result == "1":
            try:
                res = re.search(pattern=u"target='_blank' href='(.+)'", string=resp.text).group(1)  # 匹配公司详情url
            except AttributeError:
                gsxt_logger.info(f"[!]详情url获取失败!{keyword}")  # 针对 "吊销,未注销" 公司的查询
                raise NameError(u"url匹配失败!")
            self.info_url = self.info_base_url + res
            # print(self.info_url)
            return True
        elif query_result == "0":
            gsxt_logger.info(f"[!]未查询到相关信息{query_result}:{keyword}")  # 查询不到相关信息
            # print(u"查询条数", query_result)
            return False
        elif query_result:
            '''//a[@class='search_list_item db']  xpath'''
            resp = etree.HTML(resp.text)
            items = resp.xpath(u"//a[@class='search_list_item db']")
            if items:
                for item in items:
                    key = "".join([word.strip() for word in item.xpath(u".//h1//text()")])
                    if key == keyword:
                        '''//*[@id="advs"]/div/div[2]/a'''
                        url = item.xpath(u'./@href')[0]
                        self.info_url = self.info_base_url + url
                        # print(self.info_url)
                        return True
            else:
                return False
        else:
            self.query(keyword)

    # 获取公司的详细信息页面,公告信息页面,进行处理
    def get_detail_img(self, notice_url):
        '''
        使用selenium访问 url,获取公告页面和企业详情页面,并截图获取文本信息
        将图片信息作为临时文件保存在本地(./img_temp/),文本信息作为返回值
        '''
        self.date = "None"
        time.sleep(random.randint(3, 6))
        self.browser.get(notice_url)
        # 处理公告页面
        notice_wait = WebDriverWait(self.browser, 10)

        flag_notice = False
        try:
            flag_notice = notice_wait.until(
                EC.presence_of_element_located((By.ID, "datefornow"), ))
        except Exception as e:
            gsxt_logger.info(f"[!]页面请求失败:{notice_url}")
            refresh_num = 0
            while not flag_notice:
                gsxt_logger.info("[!]页面请求失败,重新刷新页面-公告")
                self.browser.get(notice_url)
                try:
                    flag_notice = notice_wait.until(
                        EC.presence_of_element_located((By.ID, "datefornow"), ))
                except Exception:
                    gsxt_logger.info(f"-刷新次数{refresh_num}:{notice_url}")
                    refresh_num += 1
                    if refresh_num == 3:
                        break
        self.notice_text = self.browser.find_element(By.XPATH,
                                                     "//div[@class='mainContent']/div[@class='contentfordetail']").text

        # 字符串 按照 \n 切割操作 取出 日期 last
        temp_data = self.notice_text.split()
        if temp_data:  # 确认公告
            if "年" in temp_data[-1]:  # 确认日期存在
                self.date = temp_data[-1]
        else:
            gsxt_logger.info(f"[!]生效时间获取失败:{notice_url}")
            # print(temp_data)
            raise NameError("AttributeError:生效时间获取失败!")  # 页面访问失败(公告栏为空)

        js = self.browser.find_element_by_xpath("//div[@class='page']/div")
        self.browser.execute_script("arguments[0].scrollIntoView(true);", js)
        self.browser.save_screenshot(f"{self.HOME_PATH}/img_temp/notice_temp.png")
        # 处理详情页面p
        time.sleep(random.randint(1, 3))
        self.browser.get(self.info_url)
        wait_info_page = WebDriverWait(self.browser, 10)
        flag = False
        try:
            flag = wait_info_page.until(
                EC.text_to_be_present_in_element(
                    (By.XPATH, "//div[@id='primaryInfo']//div[@class='classify']"),
                    '营业执照信息'))
        except Exception as e:
            gsxt_logger.info(f"[!]页面请求失败:{self.info_url}")
            refresh_num = 0
            while flag is not True and refresh_num < 3:
                gsxt_logger.info(f"[!]页面请求失败,重新刷新页面-详情")
                self.browser.get(self.info_url)
                try:
                    flag = wait_info_page.until(
                        EC.text_to_be_present_in_element(
                            (By.XPATH, "//div[@id='primaryInfo']//div[@class='classify']"),
                            '营业执照信息'))
                except Exception as e:
                    # print(f"2刷新次数{refresh_num}", e)
                    refresh_num += 1
        if self.date == "None":
            self.date = self.get_earliest_time(self.browser.page_source)
            # print("列入时间", self.date)
        else:
            pass
        js = self.browser.find_element_by_id("btn_send_pdf")
        self.browser.execute_script("arguments[0].scrollIntoView(true);", js)

        self.title = self.browser.find_element(By.XPATH, "//div[@class='overview']/dl[2]/dd[@class='result']").text
        self.reg_num = str(
            self.browser.find_element(By.XPATH, "//div[@class='overview']/dl[1]/dd[@class='result']").text)

        btn = self.browser.find_element_by_id("tab_primary5")
        btn.click()
        try:

            wait_list = WebDriverWait(self.browser, 10)
            wait_list.until(
                EC.visibility_of_element_located((By.XPATH, "//table[@id='needPaging_illegal']/tbody")))
        except:
            flag_list = False
            for i in range(0, 3):
                if flag_list:
                    break
                else:
                    # print(f"刷新次数：{i}")
                    self.browser.refresh()
                    time.sleep(2)
                    btn.click()
                    wait_list = WebDriverWait(self.browser, 10)
                    try:
                        flag_list = wait_list.until(
                            EC.visibility_of_element_located((By.XPATH, "//table[@id='needPaging_illegal']/tbody")))
                    except:
                        # print("失信列入表展示失败")
                        pass

        self.browser.save_screenshot(f"{self.HOME_PATH}/img_temp/info_temp.png")
        if len(self.reg_num) == 18:
            data_for_exc = [self.title, " ", self.reg_num, self.date]  # 公司名称,注册号,统一社会信用代码,生效日期
        else:
            data_for_exc = [self.title, self.reg_num, " ", self.date]  # 公司名称,注册号,统一社会信用代码,生效日期
        data_for_doc = {
            "notice_text": self.notice_text,  # 公告文本
            "reg_num": self.reg_num,
        }
        return data_for_exc, data_for_doc

    @staticmethod
    def check_file(date):
        if not os.path.isdir("./datas/{}".format(date)):  # todo 修改 windows环境
            os.system("mkdir ./datas//{}".format(date))
            gsxt_logger.info("[√]创建日期文件夹")
        else:
            gsxt_logger.info("[!]日期文件已经存在")
        if not os.path.isfile(f"./datas/{date}/失信企业移出信息收集.docx") or not os.path.isfile(
                f"./datas/{date}/失信名单移出登记表.xlsx"):
            os.system(
                f"cp -a {BaseConfig.HOME_PATH}/pre_data/失信企业移出信息收集.docx ./datas/{date}/失信企业移出信息收集.docx;cp -a {BaseConfig.HOME_PATH}/pre_data/失信名单移出登记表.xlsx ./datas/{date}/失信名单移出登记表.xlsx")
            gsxt_logger.info("[√]创建存储文档")
        else:
            gsxt_logger.info("[!]文档已存在")

    # 保存为excel和word
    @staticmethod
    def save_excel(data, exc_path):
        '''
        :param data:
        :param exc_path:
        :return:
        '''

        insert_columns = ["严重违法失信企业名称", "注册码", "统一社会信用代码", "生效日期"]
        df = pd.read_excel(exc_path, index_col=0, sort=False)
        update_df = pd.DataFrame([data], columns=insert_columns)
        new_data = df.append(update_df, ignore_index=True)
        new_data = new_data.drop_duplicates().reindex()
        new_data[["注册码", "统一社会信用代码"]] = new_data[["注册码", "统一社会信用代码", ]].astype('str')
        new_data.to_excel(exc_path)
        # print("保存到Excel文档成功")

    def save_doc(self, data, doc_path):
        '''
        :param data:
        :param doc_path:
        :return:
        '''
        document = Document(docx=doc_path)
        text = data.get("notice_text")
        str_list = text.split("\n")
        title = str_list[0]
        sec_title = str_list[1]
        content = "    ".join(i.strip() + "\n" for i in str_list[2:-2])
        last = "\n".join(i.strip() for i in str_list[-2:])

        # title居中
        doc_title = document.add_paragraph()
        run_title = doc_title.add_run(title)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run_title.bold = True
        # 小标题居中
        doc_title = document.add_paragraph(sec_title)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 段落
        doc_title = document.add_paragraph(content)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 落款居右
        doc_title = document.add_paragraph(last)
        doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # notice_img
        img_notice = f"{self.HOME_PATH}/img_temp/notice_temp.png"
        document.add_picture(img_notice, width=Inches(5.7))

        # 注册码
        reg_num_info = f"注册码:{data.get('reg_num')}"
        document.add_paragraph(reg_num_info)

        # info_img
        document.add_paragraph("\n")
        img_info = f"{self.HOME_PATH}/img_temp/info_temp.png"
        document.add_picture(img_info, width=Inches(5.7))
        document.add_paragraph("\n")
        document.save(doc_path)
        self.r.zincrby(self.REDIS_SET_NAME, value=str(self.brief_msg), amount=1)  # 一个redis信息处理完成,将score分数改为1,防止重复使用
        # print("保存到word文档成功")

    def __del__(self):
        # 关闭浏览器
        if self.browser:
            self.browser.quit()
        if self.display:
            self.display.stop()


def main_for_img():
    date = datetime.today().date()
    exc_path = f"/datas/shixin/{date}/失信名单移出登记表.xlsx"  # for product
    doc_path = f"/datas/shixin/{date}/失信企业移出信息收集.docx"  # for product
    exc = f"./datas/{date}/失信名单移出登记表.xlsx"
    word = f"./datas/{date}/失信企业移出信息收集.docx"
    company_info = CompanyInfo()
    company_info.check_file(date)
    company_info.open_chrome()
    error_times = 0
    global f, company_name
    f = 1
    while f:
        counts = company_info.r.zcount(BaseConfig.REDIS_SET_NAME, 0, 0)
        if counts:
            gsxt_logger.info(f"[!]待处理数据量:{counts}")
            company_info.get_cookies()
            for i in range(counts):
                try:
                    time.sleep(random.randint(10, 20))
                    notice_url, reg_num, company_name = company_info.get_redis_data()

                    if reg_num and company_info.query(keyword=reg_num):
                        # print("使用注册码查询", reg_num)
                        exc_data, doc_data = company_info.get_detail_img(notice_url)
                        company_info.save_excel(data=exc_data, exc_path=exc)
                        company_info.save_doc(data=doc_data, doc_path=word)
                    elif company_name and company_info.query(keyword=company_name):
                        # print("使用公司名称查询", company_name)
                        exc_data, doc_data = company_info.get_detail_img(notice_url)
                        company_info.save_excel(data=exc_data, exc_path=exc)
                        company_info.save_doc(data=doc_data, doc_path=word)
                    elif company_name:
                        # print(company_name)
                        data = [company_name, None, None, "未查询到详情或查询不准确"]
                        company_info.save_excel(data=data, exc_path=exc)
                        with open(BaseConfig.HOME_PATH + "/log/out_of_query.txt", 'a+', encoding="utf-8") as fp:
                            fp.write(str(company_name) + "\n")
                            company_info.r.zincrby(BaseConfig.REDIS_SET_NAME, value=str(company_info.brief_msg),
                                                   amount=2)  # 未查询到详细信息的企业,跳过,标记为2
                    else:
                        raise NameError("其他查询失败的情况")  # 个别公司查询失败的情况
                    # print(f"[√]本次运行已处理{i+1}条")
                    if error_times > 0:
                        error_times -= 1
                except EOFError:
                    error_times += 1
                    if error_times == 10:
                        break
                except NameError as e:
                    # print(e)
                    with open(BaseConfig.HOME_PATH + "/log/out_of_query.txt", 'a+', encoding="utf-8") as fp:
                        fp.write(str(company_name) + "\n")
                    company_info.r.zincrby(BaseConfig.REDIS_SET_NAME, value=str(company_info.brief_msg),
                                           amount=2)  # 公告页面未显示,标记为2
                except (Exception, BaseException):
                    with open(f"{BaseConfig.HOME_PATH}/log/error.log", "a+", encoding="utf-8") as fp:
                        fp.write(time.strftime("%Y-%m-%d %H:%M:%S",
                                               time.localtime(time.time())) + ":" + traceback.format_exc() + "\n")
        else:
            gsxt_logger.info("[!]暂无待处理数据")
            # time.sleep(60*60)
            f = 0


if __name__ == '__main__':
    try:
        main_for_img()
    except Exception as e:
        gsxt_logger.info(e)
