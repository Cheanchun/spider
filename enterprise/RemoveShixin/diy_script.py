def trans():
    r = redis.Redis(host="47.105.54.129", port=6388, db=0)
    resp = r.smembers("companys", )
    for item in resp:
        score = 0
        temp = eval(item.decode("utf-8"))
        try:
            reg_num = re.search("注册号.\d+", temp.get("noticeContent")).group()[4:]
        except Exception as e:
            score = 2
            print("注册码获取失败", score, e)
            reg_num = ""
        data = {k: v for k, v in temp.items() if v != ""}
        data["reg_num"] = reg_num
        data = str(data)
        print(data)
        result = r.zadd("order_company", {data: score})
        print(result)


def mark_score():
    r = redis.Redis(host="47.105.54.129", port=6388, db=0)
    a = r.zcount("order_company", 0, 11)
    print(a)
    while True:
        aa = r.zrange("order_company", start=0, end=1, desc=True, withscores=True, score_cast_func=int)
        print(aa)
        bb = r.zincrby("order_company", value=aa[0][0], amount=1)
        # aa = r.zrange("order_company", 2, 10, desc=False, withscores=True, score_cast_func=int)
        print(bb)


def handle_exc():
    import pandas as pd
    filename = "E:\projects\code\gest\失信名单客户登记表.xlsx"
    # 创建一个engine='openpyxl'的 ExcelWriter 对象 writer
    writer = pd.ExcelWriter(filename)
    df1 = pd.DataFrame(data={'col1': [1, 1], 'col2': [2, 2]})
    df1.to_excel(writer, 'Sheet1')
    writer.save()


import re


def pipei(string=None):
    res = re.search(pattern="([0-9A-Z]{18}[)）/])|([0-9A-Z]{13}[)）/])|([0-9A-Z]{15}[)）/])", string=string)
    if res:
        print(res, "------", res.group()[:-1])
    else:
        print("匹配失败", string)


text_string = [
    "统一社会信用码/注册号：chengchangchu1234）",  # 1 5
    "统一社会信用码/注册号：43121234432143214321mklo）",  # 1 5
    "统一社会信用码/注册号：430211000014489）",  # ok  ->15
    "统一社会信用码/注册号：123456789ABCDEF）",  # ok  ->15
    "统一社会信用码/注册号/91520521599397170L）",  #
    "统一社会信用码/注册号：91520114599368249A）",  #
    "统一社会信用码/注册号：123456789101a）",  #
    "统一社会信用码/注册号：91220622753609670A/2206221250280）",  #
    "统一社会信用码/注册号2207210000134761）",  #
    "统一社会信用码/注册号220721000013476）",  #
    r'''统一社会信用代码/注册号/220721000021601):</p><div style="margin-top: 10px"><p>''',
    "统一社会信用代码/注册号912207213078784910/220721000030409):</p><"
]


# for string in text_string:
#     pipei(string)


def get_name(string):
    res = re.search(pattern="^[^关于].+(服务部|公司|店|网吧|厅|厂|馆|中心|药房|商行|经营部|服务站|招待所)", string=string)
    if res:
        print(res, "-------", res.group())
        return
    res = re.search(pattern="关于(.+(服务部|公司|店|网吧|厅|厂|馆|中心|药房|商行|经营部|服务站|招待所))", string=string)
    if res:
        print(res, "------", res.group())
        return
    if not res:
        print("匹配失败", string)


# 以 关于......列入....   为主   极少数: 以 店 网吧 厅 结尾
test_list = [
    '关于新疆融汇信德投资管理有限公司和田分公司列入严重违法失信企业名单公告',
    '关于阿勒泰市金山砖厂列入严重违法失信企业名单公告',
    '关于新疆融汇信德投资管理有限公司和田分公司列入严重违法失信企业名单公告',
    '关于阿勒泰市金山砖厂列入严重违法失信企业名单公告',
    '关于西安华国湖南土菜馆的列入严重违法失信企业公告',
    '关于阿勒泰市金山砖厂列入严重违法失信企业名单公告',
    '关于西安市碑林区黄雁综合商店的列入严重违法失信企业公告',
    '关于西安新中阳影像制作有限公司热带鱼浪漫写真婚纱骡马市分店的列入严重违法失信企业公告',
    '关于西安裔璨商贸有限公司的列入严重违法失信企业公告',
    '福建誉衡投资有限公司列入严重违法失信企业名单',
    '福建省清禾生态农业开发有限公司列入严重违法失信企业名单',
    "关于星烨世纪（北京）国际文化传播有限公司列入严重违法失信企业名单的决定",
    "关于铸梦堂（北京）商业管理有限公司列入严重违法失信企业名单的决定",
    "关于大有盛世（北京）文化传播有限公司列入严重违法失信企业名单的决定",
    "关于东方渊博（北京）教育咨询有限公司列入严重违法失信企业名单的决定",
    "关于晓晨日升国际电控设备（北京）有限公司列入严重违法失信企业名单的决定",
    "关于北京顺华德风货运信息咨询中心列入严重违法失信企业名单的决定",
    "关于中鼎星辉（北京）国际贸易有限公司列入严重违法失信企业名单的决定",
    "关于欧佰怡（北京）公关咨询有限公司列入严重违法失信企业名单的决定",
    "关于积海（北京）国际投资管理有限公司列入严重违法失信企业名单的决定",
    "关于北京建发林业发展中心列入严重违法失信企业名单的决定",
    "关于北京鑫鑫益德不锈钢材料销售中心列入严重违法失信企业名单的决定",
    "关于北京芳华堂大药房列入严重违法失信企业名单的决定",
    "北溪（永春）白鹤拳文化传播有限责任公司列入严重违法失信企业名单",
]

# encoding: utf-8



# class Doc2:
#     def __init__(self, file_name=None):
#         self.document = Document(docx=file_name)
#         print("doc2")
#
#     def add_text(self):
#         self.document.add_paragraph("fhdahfldsa")
#         pass
#
#     def add_img(self):
#         self.document.add_picture("pic_data", width=Inches())
#         pass

import redis
from docx.shared import Inches

r = redis.Redis(host="47.105.54.129", port=6388, password="admin", db=0)


# r.zadd("test", {str("1"): 0})
# res = r.zincrby("test", value=str(str("1")), amount=2)
# resp = r.zrank("test", "2")
# brief_msg = r.zrange("order_company", start=0, end=0, desc=False, withscores=True, score_cast_func=int)[0]
# print(brief_msg)
# res = r.zincrby("order_company", value=str(brief_msg), amount=5)

# brief_msg = r.zrange("order_company", start=0, end=0, desc=True, withscores=False, score_cast_func=int)
# print(brief_msg)
# resu = r.zrange("order_company", start=0, end=0, desc=False, withscores=True, score_cast_func=int)
# print("-----", resu)
# print("+++++", resu[0][0].decode("utf-8"))
# res = r.zincrby("order_company", value=str(resu[0][0].decode("utf-8")), amount=1)


# if not resp:
#     print(res)
#     print(resp)

def prox():
    from selenium import webdriver
    chromeOptions = webdriver.ChromeOptions()

    # 设置代理
    chromeOptions.add_argument("--proxy-server=http://202.20.16.82:10152")
    # 一定要注意，=两边不能有空格，不能是这样--proxy-server = http://202.20.16.82:10152
    browser = webdriver.Chrome(chrome_options=chromeOptions)

    # 查看本机ip，查看代理是否起作用
    browser.get("http://httpbin.org/ip")
    print(browser.page_source)

    # 退出，清除浏览器缓存
    browser.quit()


from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def doc_handle():
    document = Document(docx="./失信企业移出信息收集.docx")
    title = "关于广宗县鑫诺标准件轴承厂列入严重违法失信企业名单的公告"
    sec_title = "冀工商列严字[2018]第60009号"
    text = '''
经查，广宗县鑫诺标准件轴承厂（统一社会信用码/注册号：130531200001454）因被列入经营异常名录届满3年仍未履行相关义务的,违反工商行政管理法律、行政法规且情节严重。依据《严重违法失信企业名单管理暂行办法》第五条第一款第（一）项的规定，作出决定：将上述企业列入严重违法失信企业名单。
你单位如对本决定有异议，可以自公示之日起三十日内向我局提出书面申请并提交相关证明材料，要求核实。也可在接到本决定书之日起六十日内向国家工商行政管理总局工商行政管理总局或者河北省人民政府申请行政复议；或者六个月内向石家庄市裕华区人民法院提起行政诉讼。
'''

    last = "河北省工商行政管理局\n2018年12月11日"

    # title居中
    doc_title = document.add_paragraph()
    run_title = doc_title.add_run(title)
    doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_title.bold = True
    # 小标题居中
    doc_title = document.add_paragraph(sec_title)
    doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 段落
    doc_title = document.add_paragraph(text)
    doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # 落款居右
    doc_title = document.add_paragraph(last)
    doc_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # notice_img
    img_notice = "./notice/text.png"
    document.add_picture(img_notice, width=Inches(5.7))

    # 注册码
    reg_num = 543254325426543
    reg_num_info = f"注册码:{reg_num}"
    document.add_paragraph(reg_num_info)
    # info_img
    document.add_paragraph("\n")
    img_info = "./info/temp.png"
    document.add_picture(img_info, width=Inches(5.7))
    document.add_paragraph("\n")

    print(document)
    document.save("./失信企业移出信息收集.docx")


def split():
    string = '''关于河北松立机电设备贸易有限公司列入经营异常名录的公告
石西市监异字[2018]4024号
经查，河北松立机电设备贸易有限公司（统一社会信用码/注册号：130100000363655）因通过登记的住所或者经营场所无法联系，违反了《企业信息公示暂行条例》和《企业经营异常名录管理办法》的相关规定。根据《企业信息公示暂行条例》第十七条第一款和《企业经营异常名录管理办法》第九条的规定，现决定将你单位列入经营异常名录。
　　你单位如不服本决定，可在接到本决定书之日起六十日内向石家庄市工商行政管理局或者石家庄市桥西区人民政府申请行政复议；或者六个月内向石家庄市桥西区人民法院提起行政诉讼。
桥西区市场监督管理局

2018年12月11日'''
    str_list = string.split("\n")
    title = str_list[0]
    sec_title = str_list[1]
    content = "\n".join(i.strip() for i in str_list[2:-3])
    last = "\n".join(i.strip() for i in str_list[-3:])

    print(title)
    print(sec_title)
    print(content)
    print(last)


proxies = [
    # {'http': 'http://223.145.212.45:8118'},
    # {'http': 'http://61.135.217.7:80'},
    {'http': 'http://223.145.212.45:8118'},  # 快
    # {'http': 'http://61.135.217.7:80'},  # fail
    {'http': 'h113.116.177.108:808'},  # ok
    {'http': 'http://27.155.83.182:8081'},  # 快
    {'http': 'http://219.234.5.128:3128'},  # ok
    {'http': 'http://222.182.121.84:8118'},  # ok 快
    {'http': 'http://61.135.217.7:80'},
    {'http': 'http://111.177.166.226:9999'},
    {'http': 'http://61.135.217.7:80'},
    {'http': 'http://111.177.166.92:9999'},
    {'http': 'http://219.238.186.188:8118'},
    {'http': 'http://113.116.48.171:808'},
    {'http': 'http://106.12.7.54:8118'},
    {'http': 'http://111.177.178.132:9999'},
    {'http': 'http://171.35.52.182:8118'},
    {'http': 'http://119.101.113.26:9999'},
    {'http': 'http://119.101.114.104:9999'},
    {'http': 'http://113.120.62.247:9999'},
    {'http': 'http://119.101.117.30:9999'},
    {'http': 'http://219.234.5.128:3128'},
    {'http': 'http://116.7.176.60:8118'},
    {'http': 'http://202.119.248.8:80'},
    {'http': 'http://139.129.207.72:808'},
    {'http': 'http://123.15.58.124:61202'},
]


def test_for_sele(proxy):
    import re
    from selenium import webdriver
    from selenium.webdriver.support.ui import WebDriverWait  # 等待
    url = "http://www.gsxt.gov.cn"
    option = webdriver.ChromeOptions()
    option.add_argument(f"--proxy-server={proxy.get('http')}")
    # option.headless = True
    browser = webdriver.Chrome(chrome_options=option)
    browser.get(url=url)
    WebDriverWait(browser, 10)
    resp = browser.page_source
    pattern = '代理出现问题'
    s = re.findall(pattern=pattern, string=resp)
    if s:
        print(proxy, "可用")
    print(s)


def req_test(proxy):
    '''
    测试ip可用性
    :return: None
    '''
    import requests
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/72.0.3610.2 Safari/537.36',
    }
    url = "http://www.gsxt.gov.cn"
    res = requests.get(url=url, headers=headers, proxies=proxy)
    print(res.text)


def change_redis_zset_data(min=3, max=3, desc=True, amount=0, zset_name="order_company", db=0):
    '''
    # 设置有序集合的分数
    :param min: 最小分数
    :param max: 最大分数
    :param desc: 是否倒序
    :param amount: 原有分数上的增量
    :param zset_name: 有序集合名称
    :return: None
    '''
    r = redis.Redis(host="47.105.54.129", port=6388, db=db, password="admin", )

    counts = r.zcount(zset_name, min=min, max=max)  # 统计分数起始个数
    print(counts)
    for i in range(counts):
        key = r.zrange(zset_name, start=0, end=0, desc=desc, withscores=True,
                       score_cast_func=int)
        key = key[0][0].decode("utf-8")  # start 和 end 不是分数参数,是集合元素索引的值,
        # start 和 end 不是分数参数,是集合元素索引的值, 返回值是索引区间的元素,为列表
        print(key)
        r.zincrby("order_company", value=key, amount=amount)
        # amount 在元素原有分数上进行增减  key 需要是字符串,redis查询出来的结果为字节流,需要转码才能以value的值传入才能匹配唯一性,否则会重新写入一个新值


def list_unique_dict(li: list):
    '''
    列表中的字典去重,返回去重后的字典
    :param li:
    :return: res
    '''
    res = [eval(item) for item in list(set([str(i) for i in li]))]
    return res


def get_title_to_txt():
    counts = r.zcount("order_company", min=0, max=2)
    titles = []
    for i in range(counts):
        key = r.zrange("order_company", start=i, end=i, desc=False, withscores=True, score_cast_func=int)[0][
            0].decode("utf-8")
        brief_msg = eval(key)
        title = brief_msg.get("noticeTitle")

        titles.append(title)
    with open("./title.txt", "w", encoding="utf-8") as fp:
        fp.write("\n".join(titles))


if __name__ == '__main__':
    def get_name():
        counts = r.zcount("order_company", min=0, max=2)
        titles = []
        for i in range(counts):
            key = r.zrange("order_company", start=i, end=i, desc=False, withscores=True, score_cast_func=int)[0][
                0].decode("utf-8")
            brief_msg = eval(key)
            title = brief_msg.get("noticeTitle")
            pattern = "关于(.+)(列入+?|的列入+?)"
            res = re.match(pattern=pattern, string=title)
            if not res:
                pattern = "(^[^关于].+)(列入+?|的列入+?)"
                res = re.match(pattern=pattern, string=title)
                try:
                    name = res.group(1)
                    if name[-2:] == "企业":
                        name = name[:-2]
                    elif name[-1:] == "的":
                        name = name[:-1]
                except:
                    name = ""
                    print("匹配失败", title)
            else:
                try:
                    name = res.group(1)
                    if name[-2:] == "企业":
                        name = name[:-2]
                    elif name[-1:] == "的":
                        name = name[:-1]
                except:
                    name = ""
                    print("匹配失败", title)
            if name in ["广州市荔湾区工商行政管理局", "广州市白云区工商行政管理局", "广州市海珠区工商行政管理局", "广州市增城区工商行政管理局", "广州市工商行政管理局", "辽宁省市场监督管理局"]:
                print(name)
                title = brief_msg.get("noticeContent")
                pattern_content_company = "(.+?)（统一社会信用代码/"
                result = re.match(pattern=pattern_content_company, string=title)
                try:
                    name = result.group(1)
                except:
                    name = ""
                    print("匹配失败", title)
            if name in ["吉林省工商行政管理局"]:
                print(name)
                title = brief_msg.get("noticeContent")
                result = re.findall(pattern="(<p text-align: center font-size:16px>|<p>)(.+)</p>", string=title)
                try:
                    # name = result.group().split("</p><p>")[0][3:]
                    name = result[0][1].split("</p><p>")[1]
                    if "统一社会信用代码" in name:
                        name = re.match(pattern="(.+?)[\(]", string=name).group(1)
                except:
                    name = ""
                    print("匹配失败", title)
            print(title, "----------", name)


    change_redis_zset_data(min=1, max=1, desc=True, amount=-1, zset_name="order_company", db=1)
